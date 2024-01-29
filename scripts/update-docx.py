from docx import Document
from docx.shared import Inches

document = Document("../documents/HLM.docx")

dict = {
    "NomBailleur1": "3F Construction SA",
    "NomProgramme1": "Les alouettes",
    "Type": "PLAI",
    "NbLogement1": "21",
    "Fi1": "CAF+",
    "Adresse1": "Avenue du général Toto",
    "CodePostal1": "13001",
    "Ville1": "Marseille",
}


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


for paragraph in document.paragraphs:
    if paragraph.text == "__Image__":
        paragraph.text = ""
        paragraph.add_run().add_picture(
            "../documents/Screenshot.png", width=Inches(5.0)
        )
        next  # noqa: B018
    if paragraph.text == "__Tableau__":
        paragraph.text = "Voici mon tableau:"
        records = (
            (3, "101", "Spam"),
            (7, "422", "Eggs"),
            (4, "631", "Spam, spam, eggs, and spam"),
        )
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Qty"
        hdr_cells[1].text = "Id"
        hdr_cells[2].text = "Desc"
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc
        move_table_after(table, paragraph)
        next  # noqa: B018
    text = paragraph.text
    text_updated = False
    for key in dict.keys():
        key_in_text = "__" + key + "__"
        if key_in_text in text:
            text_updated = True
            text = text.replace(key_in_text, dict[key])
    if text_updated:
        paragraph.text = text


document.save("../documents/generated_doc1.docx")


# from docxtpl import DocxTemplate

# doc = DocxTemplate("../documents/HLM.docx")
# context = { 'nombailleur' : "World company" }
# doc.render(context)
# doc.save("../documents/generated_doc2.docx")
